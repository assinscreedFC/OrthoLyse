import json
import csv
import sys

from fpdf import FPDF  # type:ignore
from docx import Document  # type:ignore
from datetime import datetime
import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL


def _normalize_data(data):
    """
    Remplace les '_' par ' ' dans les clés du dictionnaire.
    """
    return {k.replace('_', ' '): v for k, v in data.items()}


def exporte_json(chemin, data):
    """
    Exporte les données en format JSON (clés normalisées).
    """
    data = _normalize_data(data)
    with open(chemin, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


def exporte_txt(chemin, data):
    """
    Exporte les données en format texte (.txt). Clés normalisées.
    """
    data = _normalize_data(data)
    with open(chemin, "w", encoding="utf-8") as f:
        for cle, valeur in data.items():
            f.write(f"{cle}: {valeur}\n")


def exporte_csv_column(chemin, data):
    """
    Exporte les données en format CSV (clé/valeur). Clés normalisées.
    """
    data = _normalize_data(data)
    with open(chemin, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Clé", "Valeur"])
        for cle, valeur in data.items():
            writer.writerow([cle, valeur])


def exporte_csv_row(chemin, data):
    """
    Exporte les données en format CSV sous forme de tableau.
    Clés normalisées.
    """
    data = _normalize_data(data)
    with open(chemin, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(data.keys())
        values = [v if isinstance(v, list) else [v] for v in data.values()]
        for row in zip(*values):
            writer.writerow(row)


def exporte_pdf(chemin, data, titre, font_dir='../frontend/assets/Fonts/Poppins'):
    """
    Exporte les données en format PDF avec un titre, un tableau et la transcription.
    Clés normalisées.
    """
    data = _normalize_data(data)
    pdf = FPDF()
    pdf.add_page()
    font_dir = os.path.abspath(font_dir)
    pdf.add_font('Poppins-Bold', '', os.path.join(font_dir, 'Poppins-Bold.ttf'), uni=True)
    pdf.add_font('Poppins', '', os.path.join(font_dir, 'Poppins-Regular.ttf'), uni=True)
    pdf.set_font('Poppins-Bold', '', 24)
    pdf.cell(0, 12, txt=titre, ln=True, align='C')
    pdf.ln(10)
    now = datetime.now()
    date_str = now.strftime("%d/%m/%Y")
    time_str = now.strftime("%H:%M")
    pdf.set_line_width(0.8)
    padding = 20
    col_w = (pdf.w - 2 * padding) / 2
    row_h = 14
    start_x = padding

    def table_row(key, value):
        pdf.set_x(start_x)
        pdf.set_font('Poppins-Bold', '', 14)
        pdf.cell(col_w, row_h, str(key), border=1, align='C')
        pdf.set_font('Poppins', '', 14)
        pdf.cell(col_w, row_h, str(value), border=1, ln=1, align='C')

    table_row("Date", date_str)
    table_row("Heure", time_str)
    for cle, valeur in data.items():
        if cle != "texte":
            table_row(cle, valeur)

    if "texte" in data:
        pdf.ln(15)
        pdf.set_font('Poppins-Bold', '', 20)
        pdf.cell(0, 12, txt="Texte de la transcription", ln=True, align='L')
        pdf.ln(5)
        pdf.set_font('Poppins', '', 12)
        pdf.multi_cell(0, 8, txt=data.get("texte", ""))

    pdf.output(os.path.join(os.path.abspath(""), f"{titre}.pdf"))


def exporte_docx(chemin, data, titre, font_path='../frontend/assets/Fonts/Poppins-Regular.ttf'):
    """
    Exporte les données en format DOCX avec un titre, un tableau et la transcription.
    Clés normalisées.
    """
    data = _normalize_data(data)
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Poppins'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:ascii'), 'Poppins')
    style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Poppins')

    # Titre
    titre_paragraph = doc.add_paragraph()
    run = titre_paragraph.add_run(titre)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Poppins'
    run.font.color.rgb = RGBColor(0, 0, 0)
    titre_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    now = datetime.now()
    date_str = now.strftime("%d/%m/%Y")
    time_str = now.strftime("%H:%M")

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def set_cell_height(cell, h):
        tr = cell._tc.getparent()
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(h))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    def add_row(key, val):
        cells = table.add_row().cells
        left, right = cells
        run_key = left.paragraphs[0].add_run(str(key))
        run_key.font.name = 'Poppins'
        run_key.font.size = Pt(14)
        run_key.bold = True
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_height(left, 600)

        run_val = right.paragraphs[0].add_run(str(val))
        run_val.font.name = 'Poppins'
        run_val.font.size = Pt(14)
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_height(right, 600)

    add_row("Date", date_str)
    add_row("Heure", time_str)
    texte = None
    for cle, val in data.items():
        if cle == "texte":
            texte = val
        else:
            add_row(cle, val)

    if texte:
        doc.add_paragraph()
        h1 = doc.add_paragraph().add_run("Texte de la transcription")
        h1.bold = True
        h1.font.name = 'Poppins'
        h1.font.size = Pt(20)
        para = doc.add_paragraph()
        run = para.add_run(texte)
        run.font.name = 'Poppins'
        run.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(os.path.join(os.path.abspath(""), f"{titre}.docx"))


# Exemple d'utilisation
if __name__ == '__main__':
    data = {
        "nom_utilisateur": "Alice",
        "âge": 25,
        "profession": "Développeuse",
        "ville_residence": "Paris",
        "texte": "Exemple de texte."
    }
    exporte_json('output.json', data)
    exporte_txt('output.txt', data)
    exporte_csv_column('output_col.csv', data)
    exporte_csv_row('output_row.csv', data)
    exporte_pdf('', data, "rapport")
    exporte_docx('', data, "rapport")
